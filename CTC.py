import requests
import pandas as pd
import re
from collections import Counter
from datetime import datetime
import docx
import io
import ast
import time

#Insert brands
brands = ['']
#Insert Scaleserp.com API key
scaleserpkey = ''
#Insert APIFlash API key
apiflashkey = ''

for brand in brands:
    excel_df = pd.read_excel('Brand_Terms.xlsx')
    terms = excel_df.loc[:, 'Keyword']
    terms = [str(term).replace('brand', f'{brand}') for term in terms]

    contains_adverts = []
    contains_brand = []
    contains_ctc = []
    domains_in_ads = []
    domains_in_matched_brand_ads = []
    search_blobs = []


    def submit_term(term):
        params = {
            f'api_key': {scaleserpkey},
            f'q': {term},
            'output': 'json',
            'device': 'mobile',
            'gl': 'uk',
            'hl': 'en',
            'location': 'Wales,United Kingdom',
            'google_domain': 'google.co.uk'
        }

        # make the http GET request to Scale SERP
        try:
            api_result = requests.get('https://api.scaleserp.com/search', params)
        except:
            time.sleep(30)
            api_result = requests.get('https://api.scaleserp.com/search', params, headers={'Connection':'close'})

        # print the JSON response from Scale SERP
        result = api_result.json()
        return result


    def check_for_ads(serp_result):
        try:
            if serp_result['ads']:
                contains_adverts.append('yes')
        except:
            contains_adverts.append('no')


    def ctc_check(serp_result):
        try:
            search = "is_phone_ad"
            search_result = []
            for serp in serp_result['ads']:
                if search in serp.keys():
                    search_result.append(search)
            if search in search_result:
                contains_ctc.append('yes')
            else:
                contains_ctc.append('no')
        except:
            contains_ctc.append('no')


    def check_for_brand_string(serp_result):
        try:
            searchterm = [f'{brand}', f'{brand.lower()}']
            words_re = re.compile("|".join(searchterm))
            yeslist = []
            for serp in serp_result['ads']:
                string = ','.join(''.join((str(value))) for key, value in serp.items())
                stres = str(string)
                if words_re.search(stres):
                    print("----------------------------")
                    print(f"{brand} string found!")
                    print("----------------------------")
                    print("stres:", stres)
                    domains_in_matched_brand_ads.append(serp['domain'])
                    yeslist.append('yes')
            if yeslist:
                contains_brand.append('yes')
            else:
                contains_brand.append('no')
        except:
            contains_brand.append('no')


    def get_json_blob(serp_result):
        try:
            search_blobs.append(serp_result['ads'])
        except:
            search_blobs.append('')


    def count_domains(serp_result):
        try:
            for serp in serp_result['ads']:
                domains_in_ads.append(serp['domain'])
        except:
            pass


    output = {'Search Terms': terms, 'Google Ads': contains_adverts, 'Click to Call': contains_ctc,
              f'Contains "{brand}"': contains_brand, 'Raw SERP': search_blobs}

    for term in terms:
        if term:
            res = submit_term(term)
            check_for_ads(res)
            ctc_check(res)
            check_for_brand_string(res)
            count_domains(res)
            get_json_blob(res)

    current_date = datetime.now().strftime("%d-%m-%Y")

    df = pd.DataFrame(output,
                      columns=['Search Terms', 'Google Ads', 'Click to Call', f'Contains "{brand}"', 'Raw SERP'])
    df.to_excel(f"{brand}_CMC_{current_date}.xlsx", index=False)

    doc = docx.Document()
    doc.add_heading(f'CMC Google results: "{brand}"', 0)
    font = doc.styles['Normal'].font
    font.name = 'Calibri'
    font.size = docx.shared.Pt(11)

    print("Count of domains in ads: ")
    print(len(Counter(domains_in_ads)))
    print(Counter(domains_in_ads))
    count_domains_in_ads = Counter(domains_in_ads)
    print(len(count_domains_in_ads))
    r = doc.add_paragraph('Count of domains in Google Ads: ')
    alltable = doc.add_table(rows=1, cols=2)
    alltable.style = 'Light Grid Accent 1'
    hdr_cells = alltable.rows[0].cells
    hdr_cells[0].text = 'Domain'
    hdr_cells[1].text = 'Count'
    for dom, cnt in count_domains_in_ads.most_common():
        print("DOM: " + dom, "Count: " + str(cnt))
        row_cells = alltable.add_row().cells
        row_cells[0].text = dom
        row_cells[1].text = str(cnt)

    print(f"Count of domains in {brand} matched ads: ")
    print(len(Counter(domains_in_matched_brand_ads)))
    print(Counter(domains_in_matched_brand_ads))
    count_domains_in_matched_brand_ads = Counter(domains_in_matched_brand_ads)
    print(len(count_domains_in_matched_brand_ads))
    e = doc.add_paragraph('')
    e = doc.add_paragraph('Count of domains in matched Google Ads: ')
    matchtable = doc.add_table(rows=1, cols=2)
    matchtable.style = 'Light Grid Accent 1'
    hdr_cells = matchtable.rows[0].cells
    hdr_cells[0].text = 'Domain'
    hdr_cells[1].text = 'Count'
    for dom, cnt in count_domains_in_matched_brand_ads.most_common():
        print("DOM: " + dom, "Count: " + str(cnt))
        row_cells = matchtable.add_row().cells
        row_cells[0].text = dom
        row_cells[1].text = str(cnt)

    doc.add_page_break()

    for index, row in df.iterrows():
        if row[f'Contains "{brand}"'] == 'yes':
            print(row['Search Terms'])
            p = doc.add_heading('Google search term: ', level=1)
            p.add_run(str(row['Search Terms']))

            for ser in row['Raw SERP']:
                for key, value in ser.items():
                    print(key, value)
                    searchterm = [f'{brand}', f'{brand.lower()}']
                    words_re = re.compile("|".join(searchterm))
                    value = str(value)
                    if words_re.search(value):
                        found_link = (f"{brand} string found in {key}: {value}")
                        print(found_link)
                        print("..........................................")
                        print(ser['domain'])
                        domain = (f"Domain: {ser['domain']}")
                        print("..........................................")
                        print()
                        doc.add_paragraph(found_link)
                        doc.add_paragraph(domain)

                        try:
                            try:
                                if value.startswith('htt'):
                                    req = f"https://api.apiflash.com/v1/urltoimage?access_key={apiflashkey}&format=png&response_type=image&url={value}"
                                else:
                                    req = f"https://api.apiflash.com/v1/urltoimage?access_key={apiflashkey}&format=png&response_type=image&url=https://{value}"

                            except:
                                if ser['displayed_link'].startswith('htt'):
                                    req = f"https://api.apiflash.com/v1/urltoimage?access_key={apiflashkey}&format=png&response_type=image&url={ser['displayed_link']}"
                                else:
                                    req = f"https://api.apiflash.com/v1/urltoimage?access_key={apiflashkey}&format=png&response_type=image&url=https://{ser['displayed_link']}"
                            response = requests.get(req, stream=True)
                            image = io.BytesIO(response.content)
                            doc.add_picture(image, width=docx.shared.Cm(15), height=docx.shared.Cm(7.5))
                            doc.add_paragraph("")
                        except:
                            doc.add_paragraph("No image found, website not available")
    doc.save(f'{brand}_CMC_{current_date}.docx')
